"""DOCX/PDF 导出工具，支持插图。"""
from io import BytesIO
from pathlib import Path
from typing import Iterable

from .agents.base import PaperContext


def _safe_image_paths(image_paths: Iterable[str] | None) -> list[str]:
    out = []
    for p in image_paths or []:
        if not p:
            continue
        if Path(p).exists():
            out.append(p)
    return out


def export_docx_bytes(ctx: PaperContext, image_paths: Iterable[str] | None = None) -> bytes:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(ctx.topic, level=0)

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(ctx.abstract or "")
    if ctx.keywords:
        doc.add_paragraph("关键词：" + "；".join(ctx.keywords))

    doc.add_heading("引言", level=1)
    doc.add_paragraph(ctx.introduction or "")

    doc.add_heading("正文", level=1)
    for para in (ctx.content or "").split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_heading("总结与结论", level=1)
    doc.add_paragraph(ctx.conclusion or "")

    doc.add_heading("参考文献", level=1)
    for ref in ctx.references:
        doc.add_paragraph(ref)

    imgs = _safe_image_paths(image_paths)
    if imgs:
        doc.add_heading("插图", level=1)
        for idx, img in enumerate(imgs, start=1):
            doc.add_paragraph(f"图 {idx}")
            doc.add_picture(img, width=Inches(5.8))

    buff = BytesIO()
    doc.save(buff)
    return buff.getvalue()


def export_pdf_bytes(ctx: PaperContext, image_paths: Iterable[str] | None = None) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    buff = BytesIO()
    c = canvas.Canvas(buff, pagesize=A4)
    width, height = A4
    margin = 40
    y = height - margin

    font_name = "Helvetica"
    simsun = Path(r"C:\Windows\Fonts\simsun.ttc")
    if simsun.exists():
        pdfmetrics.registerFont(TTFont("SimSun", str(simsun)))
        font_name = "SimSun"
    c.setFont(font_name, 12)

    def new_page():
        nonlocal y
        c.showPage()
        c.setFont(font_name, 12)
        y = height - margin

    def draw_line(text: str, size: int = 12, gap: int = 18):
        nonlocal y
        if y < margin + 30:
            new_page()
        c.setFont(font_name, size)
        c.drawString(margin, y, text[:100])
        y -= gap

    def draw_paragraph(title: str, content: str):
        draw_line(title, size=14, gap=22)
        for raw in (content or "").split("\n"):
            line = raw.strip()
            if not line:
                y_pad()
                continue
            while line:
                chunk = line[:50]
                draw_line(chunk, size=12, gap=18)
                line = line[50:]

    def y_pad():
        nonlocal y
        y -= 8
        if y < margin + 30:
            new_page()

    draw_line(ctx.topic, size=16, gap=26)
    y_pad()
    draw_paragraph("摘要", ctx.abstract)
    if ctx.keywords:
        draw_line("关键词：" + "；".join(ctx.keywords))
    y_pad()
    draw_paragraph("引言", ctx.introduction)
    y_pad()
    draw_paragraph("正文", ctx.content)
    y_pad()
    draw_paragraph("总结与结论", ctx.conclusion)
    y_pad()
    draw_line("参考文献", size=14, gap=22)
    for ref in ctx.references:
        draw_line(ref, size=12, gap=18)

    imgs = _safe_image_paths(image_paths)
    if imgs:
        y_pad()
        draw_line("插图", size=14, gap=22)
        for idx, img in enumerate(imgs, start=1):
            if y < 260:
                new_page()
            draw_line(f"图 {idx}", size=12, gap=18)
            try:
                reader = ImageReader(img)
                c.drawImage(reader, margin, y - 220, width=width - margin * 2, height=200, preserveAspectRatio=True)
                y -= 230
            except Exception:
                draw_line(f"[图片加载失败] {img}", size=11, gap=16)

    c.save()
    return buff.getvalue()
